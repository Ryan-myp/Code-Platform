"""v15 翻译增强单测：术语表记忆（强制应用）+ 双语对照导出（md/docx）。

覆盖：
- _build_translation_prompt：术语表注入「强制术语表」规则块、空表不注入
- _load_glossary：按 user_id 隔离加载
- 术语表 CRUD 端点：增/查/删 + 归属校验
- _align_paragraphs / _build_bilingual_md / _build_bilingual_docx：段落级对照
- 导出端点：md/docx 文件生成 + 下载 URL
"""

import asyncio
import json
import sys
from pathlib import Path

import pytest

BACKEND = str(Path(__file__).resolve().parents[2] / "backend")
if BACKEND not in sys.path:
    sys.path.insert(0, BACKEND)

USER = {"user_id": "u1", "username": "user1"}
OTHER = {"user_id": "u2", "username": "user2"}


class TestTranslationPrompt:
    def test_glossary_block_injected(self):
        from extended_api import _build_translation_prompt

        glossary = [
            {"source_term": "人工智能", "target_term": "Artificial Intelligence (AI)"},
            {"source_term": "大模型", "target_term": "LLM"},
        ]
        prompt = _build_translation_prompt("中文", "English", glossary)
        assert "强制术语表" in prompt
        assert "人工智能 → Artificial Intelligence (AI)" in prompt
        assert "大模型 → LLM" in prompt
        assert "不得另行翻译" in prompt

    def test_no_glossary_no_block(self):
        from extended_api import _build_translation_prompt

        prompt = _build_translation_prompt("中文", "English")
        assert "强制术语表" not in prompt
        prompt2 = _build_translation_prompt("中文", "English", [])
        assert "强制术语表" not in prompt2

    def test_glossary_capped_at_50(self):
        from extended_api import _build_translation_prompt

        glossary = [{"source_term": f"词{i}", "target_term": f"term{i}"} for i in range(60)]
        prompt = _build_translation_prompt("中文", "English", glossary)
        assert "词49 → term49" in prompt  # 前50条生效
        assert "词59" not in prompt  # 超出50条截断


class TestGlossaryApi:
    def test_crud_and_user_isolation(self, setup_test_db):
        import extended_api
        from extended_api import GlossaryItemRequest

        # 新增
        resp = asyncio.run(
            extended_api.add_glossary_item(
                GlossaryItemRequest(source_term="人工智能", target_term="AI"), current_user=USER
            )
        )
        assert resp["ok"] is True
        gid = resp["id"]

        # 列表（自己的可见）
        items = asyncio.run(extended_api.list_glossary(current_user=USER))
        assert len(items) == 1
        assert items[0]["source_term"] == "人工智能"
        assert items[0]["target_term"] == "AI"

        # 其他用户不可见
        others = asyncio.run(extended_api.list_glossary(current_user=OTHER))
        assert others == []

        # 删除 + 归属校验（他人删除不影响自己数据）
        asyncio.run(extended_api.delete_glossary_item(gid, current_user=OTHER))
        assert len(asyncio.run(extended_api.list_glossary(current_user=USER))) == 1
        asyncio.run(extended_api.delete_glossary_item(gid, current_user=USER))
        assert asyncio.run(extended_api.list_glossary(current_user=USER)) == []

    def test_empty_terms_rejected(self, setup_test_db):
        import extended_api
        from extended_api import GlossaryItemRequest
        from fastapi import HTTPException

        with pytest.raises(HTTPException) as exc:
            asyncio.run(
                extended_api.add_glossary_item(
                    GlossaryItemRequest(source_term="  ", target_term="AI"), current_user=USER
                )
            )
        assert exc.value.status_code == 400


class TestBilingualExport:
    def test_align_same_line_count(self):
        from extended_api import _align_paragraphs

        pairs = _align_paragraphs("第一段\n第二段", "Para 1\nPara 2")
        assert len(pairs) == 2
        assert pairs[0] == {"source": "第一段", "translation": "Para 1"}
        assert pairs[1] == {"source": "第二段", "translation": "Para 2"}

    def test_align_mismatch_falls_back_to_block(self):
        from extended_api import _align_paragraphs

        pairs = _align_paragraphs("第一段\n第二段\n第三段", "整块译文")
        assert len(pairs) == 1
        assert pairs[0]["source"].startswith("第一段")
        assert pairs[0]["translation"] == "整块译文"

    def test_build_md_content(self):
        from extended_api import _build_bilingual_md

        md = _build_bilingual_md("第一段\n第二段", "Para 1\nPara 2")
        assert md.startswith("# 双语对照翻译")
        assert "## 第 1 段" in md
        assert "**原文**" in md and "**译文**" in md
        assert "第一段" in md and "Para 1" in md

    def test_build_docx_readable(self, tmp_path):
        from extended_api import _build_bilingual_docx
        from docx import Document

        path = str(tmp_path / "out.docx")
        _build_bilingual_docx("第一段\n第二段", "Para 1\nPara 2", path)
        doc = Document(path)
        texts = [p.text for p in doc.paragraphs]
        assert any("双语对照翻译" in t for t in texts)
        assert "第一段" in texts and "Para 1" in texts

    def test_export_endpoint_md_and_docx(self, setup_test_db, monkeypatch, tmp_path):
        import extended_api
        from extended_api import TranslationExportRequest

        monkeypatch.setattr(extended_api, "TRANSLATION_EXPORT_DIR", str(tmp_path))

        md_resp = asyncio.run(
            extended_api.export_translation(
                TranslationExportRequest(source="你好", translation="Hello", format="md"),
                current_user=USER,
            )
        )
        assert md_resp["ok"] is True
        assert md_resp["download_url"].startswith("/api/translation/download/")
        assert (tmp_path / md_resp["filename"]).exists()
        content = (tmp_path / md_resp["filename"]).read_text(encoding="utf-8")
        assert "你好" in content and "Hello" in content

        docx_resp = asyncio.run(
            extended_api.export_translation(
                TranslationExportRequest(source="你好", translation="Hello", format="docx"),
                current_user=USER,
            )
        )
        assert docx_resp["filename"].endswith(".docx")
        from docx import Document

        doc = Document(str(tmp_path / docx_resp["filename"]))
        assert any("你好" in p.text for p in doc.paragraphs)

    def test_export_rejects_bad_format(self, setup_test_db):
        import extended_api
        from extended_api import TranslationExportRequest
        from fastapi import HTTPException

        with pytest.raises(HTTPException) as exc:
            asyncio.run(
                extended_api.export_translation(
                    TranslationExportRequest(source="a", translation="b", format="pdf"),
                    current_user=USER,
                )
            )
        assert exc.value.status_code == 400
